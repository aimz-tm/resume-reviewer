# app.py
# Streamlit + Groq LLM Resume Reviewer
# ------------------------------------
# Features:
# - Upload PDF or paste resume text
# - Provide target job role and optional job description
# - LLM-driven structured review (section-wise feedback, scores, missing keywords)
# - Generates a tailored "improved resume" draft
# - Download improved resume as .docx
# - All processing in-memory (no persistent storage)

import os
import io
import json
import re
from datetime import datetime
from typing import Dict, Any, Optional, Tuple, List

import streamlit as st

# PDF parsing
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

# Groq client
try:
    from groq import Groq
except Exception:
    Groq = None

# DOCX export
try:
    from docx import Document
    from docx.shared import Pt, Inches
except Exception:
    Document = None

# -------------------------
# Utilities
# -------------------------
def read_pdf_bytes(pdf_bytes: bytes) -> str:
    if fitz is None:
        raise RuntimeError("PyMuPDF (fitz) not installed. Please install 'pymupdf'.")
    text_parts = []
    with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text("text"))
    return "\n".join(text_parts).strip()

def clean_text(s: str) -> str:
    # Normalize whitespace, remove odd control chars
    s = s.replace("\x00", " ").replace("\r", "\n")
    s = re.sub(r"\n{3,}", "\n\n", s)
    s = re.sub(r"[ \t]+", " ", s)
    return s.strip()

def load_env_api_key() -> Optional[str]:
    # Prioritize Streamlit secrets if present, then env var
    key = st.secrets.get("GROQ_API_KEY", None) if hasattr(st, "secrets") else None
    return key or os.getenv("GROQ_API_KEY")

def get_groq_client() -> Groq:
    if Groq is None:
        raise RuntimeError("Groq SDK not installed. Please install 'groq'.")
    api_key = load_env_api_key()
    if not api_key:
        raise RuntimeError(
            "GROQ_API_KEY not found. Set it in environment or .streamlit/secrets.toml."
        )
    return Groq(api_key=api_key)

def default_model_name() -> str:
    # Safe defaults—adjust to models enabled for your Groq account
    # Common choices: "llama-3.1-70b-versatile", "mixtral-8x7b-32768", "llama3-70b-8192"
    return "llama-3.1-70b-versatile"

def build_system_prompt() -> str:
    return (
        "You are a meticulous resume reviewer for job seekers.\n"
        "Goals:\n"
        "- Compare the resume to the target job role and optional job description.\n"
        "- Identify missing skills/keywords, vague or redundant language, formatting and clarity improvements.\n"
        "- Provide section-wise feedback (Experience, Education, Skills, Projects, Certifications, Summary, Other).\n"
        "- Produce objective, constructive feedback with examples and rewrites.\n"
        "- Provide concise scores (0-10) for key dimensions.\n"
        "- Draft an improved version of the resume text tailored to the role while preserving factual accuracy.\n"
        "Output STRICTLY in JSON with the schema below.\n"
        "Be honest about uncertainties; do not hallucinate credentials or employers. If unsure, say so."
    )

def build_user_prompt(resume_text: str, job_role: str, job_description: str, region: str, language: str) -> str:
    # Ask for explicit JSON so we can parse reliably
    schema = {
        "role_alignment_summary": "1-2 paragraphs summarizing overall alignment.",
        "section_feedback": {
            "Experience": ["bullet feedback items..."],
            "Education": ["..."],
            "Skills": ["..."],
            "Projects": ["..."],
            "Certifications": ["..."],
            "Summary": ["..."],
            "Other": ["..."]
        },
        "missing_keywords": ["list", "of", "missing", "keywords"],
        "rewrite_suggestions": [
            {
                "original": "original snippet from resume",
                "improved": "improved rewrite tailored to role",
                "reason": "why this is better"
            }
        ],
        "scores": {
            "overall": 0,
            "alignment_to_role": 0,
            "clarity": 0,
            "impactfulness": 0,
            "keyword_coverage": 0,
            "formatting": 0
        },
        "improved_resume": "A full improved resume text in Markdown, with clear section headings.",
        "disclaimers": [
            "brief notes about assumptions, uncertainties, or ethical cautions"
        ]
    }

    prompt = f"""
Task:
- Review the resume against the target job role and the optional job description.
- Consider {region} norms if relevant. Write in {language}.
- Provide concrete, actionable, non-generic feedback with examples.
- Do not invent facts. If information is missing, suggest how to add it.

Inputs:
- Target Job Role: {job_role or "(not provided)"}
- Job Description (if any): {job_description or "(not provided)"}
- Resume:
\"\"\"{resume_text}\"\"\"

Output:
Return ONLY valid JSON that strictly matches this schema (keys required, types must match).
If a section is not applicable, return an empty list for that section.

Schema (example types):
{json.dumps(schema, indent=2)}
"""
    return prompt

def call_groq_review(resume_text: str, job_role: str, job_description: str, model_name: str, region: str, language: str, temperature: float = 0.2, max_tokens: int = 3000) -> Dict[str, Any]:
    client = get_groq_client()
    # Groq chat.completions-style call
    resp = client.chat.completions.create(
        model=model_name,
        temperature=temperature,
        max_tokens=max_tokens,
        messages=[
            {"role": "system", "content": build_system_prompt()},
            {"role": "user", "content": build_user_prompt(resume_text, job_role, job_description, region, language)},
        ],
    )
    content = resp.choices[0].message.content
    # Try to extract JSON safely
    content = content.strip()
    # handle accidental code fences
    if content.startswith("```"):
        content = re.sub(r"^```(?:json)?\s*", "", content)
        content = re.sub(r"\s*```$", "", content)
    try:
        data = json.loads(content)
        return data
    except json.JSONDecodeError:
        # fallback attempt: find first {...} block
        match = re.search(r"\{.*\}\s*$", content, flags=re.DOTALL)
        if match:
            try:
                return json.loads(match.group(0))
            except Exception:
                pass
        raise ValueError("Model did not return valid JSON. Raw content:\n\n" + content)

def make_docx_from_markdown(markdown_text: str) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx not installed. Please install 'python-docx'.")
    # Very lightweight Markdown to DOCX (headings + paragraphs + bullets)
    doc = Document()

    def add_heading_if_match(line: str) -> bool:
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if not m:
            return False
        level = len(m.group(1))
        text = m.group(2).strip()
        # Map h1..h6 roughly to Word heading levels 0..5
        doc.add_heading(text, level=level-1 if level <= 6 else 0)
        return True

    def add_bullet_if_match(line: str) -> bool:
        m = re.match(r"^\s*[-*+]\s+(.*)$", line)
        if not m:
            return False
        doc.add_paragraph(m.group(1).strip(), style="List Bullet")
        return True

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        if add_heading_if_match(line):
            continue
        if add_bullet_if_match(line):
            continue
        # Plain paragraph
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.size = Pt(11)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def summarize_missing_keywords(missing: List[str]) -> str:
    if not missing:
        return "Great news: No critical keywords appear to be missing for this role."
    return "Potentially missing or under-emphasized keywords:\n- " + "\n- ".join(sorted(set(missing), key=str.lower))

def render_scores(scores: Dict[str, Any]) -> None:
    cols = st.columns(6)
    keys = ["overall", "alignment_to_role", "clarity", "impactfulness", "keyword_coverage", "formatting"]
    for i, k in enumerate(keys):
        with cols[i]:
            v = scores.get(k, 0)
            st.metric(k.replace("_", " ").title(), f"{v}/10")

def safe_get_section_list(d: Dict[str, Any], name: str) -> List[str]:
    x = d.get("section_feedback", {}).get(name, [])
    return x if isinstance(x, list) else []

def enforce_privacy_note():
    st.caption("🔒 **Privacy:** Your files and text are processed in-memory for this session only and are not stored on the server. "
               "Clear the page to remove data from memory.")

# -------------------------
# Streamlit App
# -------------------------
st.set_page_config(page_title="LLM Resume Reviewer (Groq)", page_icon="🧠", layout="wide")
st.title("🧠 LLM Resume Reviewer")
st.write("Upload your resume (PDF or text), choose a job role, and get structured, actionable feedback powered by a Groq-hosted LLM.")

with st.sidebar:
    st.header("⚙️ Settings")
    model = st.text_input("Groq Model", value=default_model_name(), help="Example: llama-3.1-70b-versatile, mixtral-8x7b-32768, llama3-70b-8192")
    language = st.selectbox("Output Language", ["English", "Hindi", "Spanish", "French", "German", "Italian", "Portuguese"], index=0)
    region = st.selectbox("Region Norms (optional)", ["Global", "India", "US", "EU", "UK"], index=1)
    temperature = st.slider("Creativity (temperature)", 0.0, 1.0, 0.2, 0.05)
    max_tokens = st.slider("Max tokens", 512, 8192, 3000, 64)
    st.divider()
    enforce_privacy_note()
    st.caption("Tip: Provide a job description for best results.")

tab_input, tab_feedback, tab_improved = st.tabs(["1) Input", "2) Feedback", "3) Improved Resume"])

with tab_input:
    st.subheader("Resume Input")
    uploaded_file = st.file_uploader("Upload resume (PDF or .txt)", type=["pdf", "txt"])
    resume_text_area = st.text_area("Or paste resume text", height=220, placeholder="Paste your resume text here...")

    st.subheader("Target Job")
    job_role = st.text_input("Target Job Role (e.g., Data Scientist, Product Manager)")
    job_description = st.text_area("Optional: Paste Job Description", height=180, placeholder="Paste the JD here for more tailored feedback...")

    colA, colB = st.columns(2)
    with colA:
        run_btn = st.button("🚀 Review Resume", use_container_width=True, type="primary")
    with colB:
        clear_btn = st.button("🧹 Clear", use_container_width=True)

    if clear_btn:
        st.experimental_rerun()

    # Prepare resume text
    resume_text_final = ""
    if uploaded_file is not None:
        if uploaded_file.type == "application/pdf":
            try:
                resume_text_final = read_pdf_bytes(uploaded_file.read())
                st.success("Parsed PDF successfully.")
            except Exception as e:
                st.error(f"Failed to parse PDF: {e}")
        else:
            try:
                resume_text_final = uploaded_file.read().decode("utf-8", errors="ignore")
                st.success("Read text file successfully.")
            except Exception as e:
                st.error(f"Failed to read text file: {e}")

    if resume_text_area.strip():
        # Prefer pasted text if present
        resume_text_final = resume_text_area

    resume_text_final = clean_text(resume_text_final)

    if run_btn:
        if not resume_text_final:
            st.error("Please upload a resume or paste the resume text.")
        elif not model.strip():
            st.error("Please specify a Groq model.")
        else:
            with st.spinner("Thinking with the LLM..."):
                try:
                    data = call_groq_review(
                        resume_text=resume_text_final,
                        job_role=job_role.strip(),
                        job_description=job_description.strip(),
                        model_name=model.strip(),
                        region=region,
                        language=language,
                        temperature=temperature,
                        max_tokens=max_tokens,
                    )
                    st.session_state["review_data"] = data
                    st.session_state["resume_text"] = resume_text_final
                    st.session_state["job_role"] = job_role.strip()
                    st.session_state["job_description"] = job_description.strip()
                    st.success("Review complete! See the Feedback and Improved Resume tabs.")
                except Exception as e:
                    st.error(f"Model error: {e}")

with tab_feedback:
    st.subheader("Structured Feedback")
    data = st.session_state.get("review_data")
    if not data:
        st.info("Run a review first in the Input tab.")
    else:
        # Summary
        st.markdown("### 📌 Role Alignment Summary")
        st.write(data.get("role_alignment_summary", "—"))

        # Scores
        st.markdown("### 🧮 Scores")
        render_scores(data.get("scores", {}))

        # Missing Keywords
        st.markdown("### 🧩 Missing / Under-Emphasized Keywords")
        st.write(summarize_missing_keywords(data.get("missing_keywords", [])))

        # Section-wise feedback
        st.markdown("### 📚 Section-wise Feedback")
        sections = ["Summary", "Experience", "Projects", "Skills", "Education", "Certifications", "Other"]
        for sec in sections:
            items = safe_get_section_list(data, sec)
            with st.expander(f"{sec} ({len(items)} items)"):
                if items:
                    for i, item in enumerate(items, 1):
                        st.markdown(f"- **{i}.** {item}")
                else:
                    st.caption("No specific feedback.")

        # Rewrites
        st.markdown("### ✍️ Targeted Rewrite Suggestions")
        rewrites = data.get("rewrite_suggestions", [])
        if isinstance(rewrites, list) and rewrites:
            for rw in rewrites:
                with st.container():
                    st.markdown("**Original:**")
                    st.code(rw.get("original", ""), language="text")
                    st.markdown("**Improved:**")
                    st.code(rw.get("improved", ""), language="text")
                    st.caption("Why better: " + (rw.get("reason", "—")))
                    st.divider()
        else:
            st.caption("No targeted rewrite suggestions returned.")

        # Disclaimers
        dis = data.get("disclaimers", [])
        if dis:
            st.markdown("### ⚖️ Notes & Disclaimers")
            for d in dis:
                st.write(f"- {d}")

with tab_improved:
    st.subheader("Improved Resume Draft")
    data = st.session_state.get("review_data")
    if not data:
        st.info("Run a review first in the Input tab.")
    else:
        improved_md = data.get("improved_resume", "")
        if improved_md:
            st.markdown(improved_md)
            # DOCX download
            try:
                docx_bytes = make_docx_from_markdown(improved_md)
                st.download_button(
                    label="⬇️ Download Improved Resume (.docx)",
                    data=docx_bytes,
                    file_name=f"improved_resume_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            except Exception as e:
                st.warning(f"Could not generate DOCX automatically: {e}")
        else:
            st.caption("Model did not return an improved resume body.")
